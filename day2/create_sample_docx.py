from docx import Document
from pathlib import Path

base_dir = Path(__file__).resolve().parent
doc = Document()
doc.add_heading('Company AI Policy', 0)
doc.add_paragraph('This document outlines our company policy on AI tool usage.')
doc.add_paragraph('All AI-generated content must be reviewed by a human before publication.')
doc.add_paragraph('Employees must not share confidential data with external AI services.')
doc.add_paragraph('Approved AI tools include internal chatbots and summarization assistants.')
doc.add_paragraph('Violations of this policy will result in disciplinary action.')
doc.save(base_dir / 'sample.docx')
print('sample.docx created successfully')
